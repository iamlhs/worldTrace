# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r'd:\华东师大计科\6月事件文件\新系统\file_source\事件分析模块_模型与数据集说明.docx'

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3.0); s.right_margin=Cm(2.5)

sty = doc.styles['Normal']; sty.font.name='宋体'; sty.font.size=Pt(10.5)
sty.paragraph_format.line_spacing=1.5; sty.paragraph_format.space_after=Pt(6)

def h1(t):
    h=doc.add_heading(t,level=1)
    for r in h.runs: r.font.name='黑体'; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0x1a,0x1a,0x1a)
def h2(t):
    h=doc.add_heading(t,level=2)
    for r in h.runs: r.font.name='黑体'; r.font.size=Pt(15)
def h3(t):
    h=doc.add_heading(t,level=3)
    for r in h.runs: r.font.name='黑体'; r.font.size=Pt(13)

def bp(text, indent=True):
    p=doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent=Cm(0.74)
    r=p.add_run(text); r.font.name='宋体'; r.font.size=Pt(10.5)

def bu(text):
    p=doc.add_paragraph(style='List Bullet'); p.clear()
    r=p.add_run(text); r.font.name='宋体'; r.font.size=Pt(10.5)

def add_table(headers, rows_data):
    t=doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.style='Table Grid'
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=''
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.name='宋体'; r.font.size=Pt(9)
    for i,row in enumerate(rows_data):
        for j,val in enumerate(row):
            c=t.rows[i+1].cells[j]; c.text=''
            r=c.paragraphs[0].add_run(val); r.font.name='宋体'; r.font.size=Pt(9)
    doc.add_paragraph()

# ===== COVER =====
for _ in range(6): doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('事件分析模块'); r.bold=True; r.font.size=Pt(28); r.font.name='黑体'
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('模型架构与数据集说明'); r.bold=True; r.font.size=Pt(24); r.font.name='黑体'
doc.add_paragraph()
for line in ['版本 V1.0', '2026 年 6 月']:
    l=doc.add_paragraph(); l.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=l.add_run(line); r.font.size=Pt(12); r.font.name='宋体'; r.font.color.rgb=RGBColor(0x66,0x66,0x66)
doc.add_page_break()

# ===== 1. 概述 =====
h1('1  概述')
bp('事件分析模块是 EDAS（Event Detection and Analysis from Social Media）系统的核心前端分析引擎。系统基于从 Twitter 采集的约 95.5 万条推文数据，通过离线预处理阶段的深度学习模型完成事件检测、摘要生成、实体抽取和图谱构建，在前端提供交互式的事件空间可视化、目标维度分析和组织关联关系图展示。')
bp('本文档基于"事件分析模块技术架构文档"的实际技术描述，系统梳理了事件分析模块所依赖的深度学习模型、训练数据集及其与各功能点的对应关系。')

h1('2  深度学习模型总览')
bp('事件分析模块的底层依赖三个核心深度学习模型和一个传统信号处理算法，均在离线预处理阶段完成推理：')

add_table(
    ['模型/算法', '框架', '核心架构', '推理阶段', '对应功能'],
    [
        ['BART', 'HuggingFace Transformers 2.8.0', 'Encoder-Decoder Transformer\n(双向编码器 + 自回归解码器)', '离线预处理', '事件摘要生成\n关键词权重提取'],
        ['TF 事件分类模型', 'TensorFlow 1.13.0', 'CNN/RNN 文本分类\n(Embedding→Conv1D/LSTM→Softmax)', '离线预处理\n(静态图推理)', '事件类型分类\n事件等级/突发判定'],
        ['BiLSTM-CRF', 'PyTorch 1.0.0 + pytorch-crf 0.7.0', '词嵌入→BiLSTM编码→CRF解码\n(双向LSTM + 条件随机场)', '离线预处理', '命名实体识别\n(人物/地点/组织)'],
        ['峰值检测算法', '自定义 (PeakDetector)', '基于信号处理的时序极值检测\n(局部最大值 + 最小间距约束)', '离线预处理', '突发事件识别\n(Bursty Flag)'],
    ],
)

doc.add_page_break()

# ===== 3. 各模型详细说明 =====
h1('3  各模型详细说明')

# --- BART ---
h2('3.1  BART 事件摘要生成模型')
h3('3.1.1  模型简介')
bp('BART（Bidirectional and Auto-Regressive Transformer）是 Facebook/Meta 提出的序列到序列（Seq2Seq）预训练模型。不同于 BERT 的"编码器专用"架构，BART 同时具有双向编码器（Encoder）和自回归解码器（Decoder），能够执行文本生成任务——这正是事件摘要生成所需的核心能力。')
h3('3.1.2  架构与机制')
bp('BART 的编码器采用与 BERT 相同的双向自注意力机制（Bi-directional Self-Attention），对输入的推文集合进行全局上下文编码。其核心计算为缩放点积注意力（Scaled Dot-Product Attention），通过 Query、Key、Value 三个矩阵的线性变换和多头并行计算，捕获文本中词与词之间的语义依赖关系。编码器输出包含每个 token 的上下文感知隐状态表示。')
bp('BART 的解码器采用自回归生成策略：通过掩码自注意力（Masked Self-Attention）确保每个生成步骤仅能关注已输出的上文 token；通过交叉注意力（Cross-Attention）将解码器的每一层与编码器的最终输出进行交互，使生成过程能够充分利用输入文本的语义信息。解码器逐词生成摘要文本，直至输出终止符。')
h3('3.1.3  注意力权重的可解释性利用')
bp('BART 编码器最后一层的自注意力权重矩阵反映了模型在处理输入文本时对各 token 的关注程度。系统提取该权重矩阵，对每个事件摘要文本中的 token 按注意力权重排序，选取 Top-12 高权重关键词。这些关键词不仅用于事件内容的关键信息概括，其权重分布也直观展示了模型判断事件核心语义的"关注焦点"——权重越高的词，在模型对事件的理解中越关键。这一机制为分析人员提供了模型决策的可解释性依据。')
h3('3.1.4  训练/推理说明')
bp('BART 模型以 HuggingFace 预训练权重（bart-large-cnn）为基础，在领域推文数据上进行微调。推理阶段为离线批量处理：将同一时间窗口内与特定事件相关的推文拼接为输入序列（截断至 1024 tokens），经 BART 编码-解码后输出流畅的自然语言摘要。每个检测到的事件均携带一条由 BART 生成的摘要文本。')

doc.add_page_break()

# --- TF ---
h2('3.2  TensorFlow 事件分类模型')
h3('3.2.1  模型简介')
bp('系统采用基于 TensorFlow 1.13.0 框架的文本分类模型进行事件类型的自动判定。该模型以静态图（Static Graph）模式加载预训练权重，通过 `tf.get_default_graph()` 获取计算图并在 Session 中执行推理。')
h3('3.2.2  架构推测')
bp('由于模型权重文件（.pb 或 checkpoint）为编译后产物，具体架构无法直接确认。根据 TensorFlow 1.x 时代的典型文本分类实践和系统依赖的 Keras 2.2.4 版本，模型架构推测为：词嵌入层（Embedding）→ 卷积层（Conv1D）或循环层（LSTM/GRU）→ 全连接层（Dense）→ Softmax 输出层。输出为事件属于各类别的概率分布。')
h3('3.2.3  在系统中的实际应用')
bp('TF 模型在离线预处理阶段完成对每个事件的类型推断。输入为事件关联推文集合的文本特征，输出为事件类别标签。分类结果映射到系统前端的四类事件类型显示：抗议、冲突、军事活动、其他。同时，模型的输出概率可用于辅助事件等级的判定——高置信度的"冲突"类别结合峰值检测的突发概率，共同决定事件的重大等级。')

# --- BiLSTM-CRF ---
h2('3.3  BiLSTM-CRF 命名实体识别模型')
h3('3.3.1  模型简介')
bp('系统采用基于 PyTorch 1.0.0 框架的 BiLSTM-CRF 深度序列标注模型，从推文文本中自动抽取命名实体。该模型结合了双向长短期记忆网络（BiLSTM）的上下文建模能力和条件随机场（CRF）的全局标签优化能力。')
h3('3.3.2  模型结构')
bp('模型由三层组成：（1）词嵌入层——将输入文本中的每个词映射为稠密向量；（2）BiLSTM 编码层——前向 LSTM 和后向 LSTM 分别从左到右和从右到左扫描序列，将两个方向的隐状态拼接，使每个位置的表示包含完整的上下文信息；（3）CRF 解码层——在 BiLSTM 输出的发射概率矩阵基础上，学习标签间的转移约束（如 I-Person 不能直接跟在 B-Location 之后），通过维特比算法求解全局最优标签序列。')
h3('3.3.3  在系统中的实际应用')
bp('BiLSTM-CRF 模型在离线预处理阶段对推文进行命名实体识别，抽取人物（Person）、地点（Location）、组织（Organization）等实体。这些实体直接服务于：')
bu('知识图谱节点构建——人物实体成为图谱中的人物节点，地点实体映射至空间维度坐标；')
bu('实体关系挖掘——同一事件中出现的多个实体之间形成共现关系，经关系抽取后成为图谱的边；')
bu('GEXF/JSON 图谱数据文件的生成——实体和关系以结构化格式持久化存储。')

doc.add_page_break()

# --- Peak Detection ---
h2('3.4  峰值检测算法')
h3('3.4.1  算法简介')
bp('峰值检测（Peak Detection）是系统中唯一不基于深度学习的分析算法，属于信号处理领域的方法。算法功能为：在一维时间序列数据（各时间窗口内的推文数量）中检测局部最大值（峰值），从而判断某一时刻是否发生突发性事件。')
h3('3.4.2  算法原理')
bp('算法扫描推文数量的时间序列，通过预设的"峰值最小间距"（Peak Distance）参数约束，确保检测到的峰值之间保持足够的时间间隔，避免将同一事件的波动误判为多个突发。检测到的峰值时刻对应的事件被标记为"突发（Bursty）"，该标志传递至前端作为事件紧急程度的量化参考。')

doc.add_page_break()

# ===== 4. 数据集 =====
h1('4  数据集说明')
h2('4.1  数据来源与规模')
bp('EDAS 系统基于 Twitter 社交媒体平台采集的真实推文数据，覆盖香港、伊朗两个地缘热点区域。')
add_table(
    ['数据集属性', '说明'],
    [
        ['数据来源', 'Twitter API 采集的公开推文'],
        ['推文总数', '约 955,012 条'],
        ['覆盖地区', 'HongKong（香港）、Iran（伊朗）'],
        ['时间跨度', '2019–2020 年'],
        ['检测事件数量', '474 个'],
        ['语言', '英文为主，部分中文（jieba/pkuseg 分词）'],
    ],
)
h2('4.2  训练/评测数据与功能点对应关系')
add_table(
    ['功能点', '依赖模型', '训练数据', '评测方式', '输出结果'],
    [
        ['事件摘要生成', 'BART', 'CNN/DailyMail 预训练\n+ 领域推文微调', 'ROUGE 评分\n(ROUGE-1/2/L)', '自然语言事件摘要文本'],
        ['事件类型分类', 'TF 分类模型', '标注事件类型的\n推文数据集', '分类准确率\n(Accuracy)', '四分类标签\n(抗议/冲突/军事活动/其他)'],
        ['命名实体识别', 'BiLSTM-CRF', '标注实体标签的\n推文序列数据\n(BIO 标注模式)', '实体级 F1 分数\n(Token-level F1)', '实体列表\n(人物/地点/组织)'],
        ['突发事件识别', '峰值检测算法', '无（非学习算法）', '查准率/查全率\n(Precision/Recall)', '突发标志\n(Bursty True/False)'],
        ['知识图谱构建', 'BiLSTM-CRF\n+ 关系规则', 'NER 输出\n+ 共现分析', '图谱完整性校验\n(节点/边覆盖度)', 'JSON/GEXF 图谱文件'],
        ['关键词权重视觉化', 'BART 注意力权重', '无需额外训练', '与人工标注的\n一致性检验', 'Top-12 关键词\n+ 权重条形图'],
        ['因果链推理', '无（规则算法）', '无需训练', '推理合理性\n人工评审', '有向因果边\n(共享词≥4/重叠≥30%)'],
    ],
)

doc.add_page_break()

# ===== 5. 功能-模型-数据对照总表 =====
h1('5  功能点与模型对应关系总表')
bp('下表以系统三大功能模块为纵轴，列出每个功能点所依赖的模型/算法、输入数据和对外输出：')

add_table(
    ['功能模块', '子功能', '依赖模型/算法', '输入数据', '输出'],
    [
        ['事件空间\n维度分析', '地图事件定位\n与区域分色', '无（坐标数据）', '事件经纬度\n+ 区域归属', '地图可视化'],
        ['事件空间\n维度分析', '多尺度空间聚类', 'Supercluster\n层次聚类算法', '事件坐标集合\n+ 缩放级别', '聚类圆 + 事件列表'],
        ['事件目标\n维度分析', '事件类型分类', 'TF 文本分类模型', '推文文本特征', '四分类标签'],
        ['事件目标\n维度分析', '事件等级判定', 'TF 模型 + 峰值检测', '分类概率\n+ 突发概率', '四级等级标签'],
        ['事件目标\n维度分析', '事件摘要展示', 'BART 摘要生成', '推文集合', '自然语言摘要'],
        ['事件目标\n维度分析', '关键词权重分析', 'BART 注意力权重', '摘要文本', 'Top-12 关键词\n+ 权重条形图+词云'],
        ['事件目标\n维度分析', '突发标志识别', '峰值检测算法', '推文时序数据', 'Bursty True/False'],
        ['事件目标\n维度分析', '因果链推理', '共享关键词\n有向图算法', '事件关键词集合', '有向因果边'],
        ['事件目标\n组织关联\n关系图', '人物关系链\n(103/114节点)', 'BiLSTM-CRF\n+ 共现关系抽取', '推文实体抽取结果', '人物链 JSON\n+ D3 力导向图'],
        ['事件目标\n组织关联\n关系图', '卫星任务链\n(44节点)', '预置静态数据\n(MongoDB)', 'KG_chain.json', '任务链 JSON\n+ D3 力导向图'],
        ['事件目标\n组织关联\n关系图', '图谱渲染信息\n(节点数/耗时)', '力模型性能计时\n(D3 forceSimulation)', '图谱数据\n+ 渲染性能日志', '节点数 + 渲染耗时'],
    ],
)

doc.add_page_break()

# ===== 6. 修正说明 =====
h1('6  对原始描述的修正')
bp('根据"事件分析模块技术架构文档"的实际技术内容，对初始描述中的以下关键点进行修正：')

add_table(
    ['原始描述（有误）', '修正后描述', '修正依据'],
    [
        ['"基于注意力机制的BERT预训练模型，通过多头自注意力机制捕捉事件摘要文本中的双向上下文语义，用于事件类型的自动分类"',
         '事件分类由 TensorFlow 文本分类模型（CNN/RNN 架构）完成，而非 BERT。BERT 的双向自注意力机制存在于 BART 的编码器中，用于事件的摘要生成，而非分类。',
         '技术架构文档 §4：TF 模型用于事件检测分类；§3：BART 编码器含双向 Self-Attention，用于摘要生成。'],
        ['"(BERT)注意力权重可解释模型对关键词的关注程度"',
         '关键词权重来源于 BART 编码器的自注意力矩阵，而非独立 BERT 模型。BART 的 Encoder 在架构上与 BERT 同源（均为双向 Transformer 编码器），其注意力权重同样具有可解释性价值。',
         '技术架构文档 §3.4：BART 使用 Self-Attention（编码器）、Masked Self-Attention（解码器）和 Cross-Attention 三种注意力机制。'],
        ['未提及命名实体识别',
         '事件分析模块的知识图谱节点构建依赖于 BiLSTM-CRF 命名实体识别模型，该模型基于 PyTorch 框架，从推文中抽取人物、地点、组织等实体，直接影响图谱质量和目标维度分析。',
         '技术架构文档 §5：PyTorch + CRF 命名实体识别，识别 Person/Location/Organization 用于构建知识图谱节点。'],
        ['未提及数据集规模',
         '系统基于约 95.5 万条真实 Twitter 推文数据（覆盖香港和伊朗，2019-2020 年），检测出 474 个事件。知识图谱最大包含 5,287 个节点和 4,191 条边。',
         '技术架构文档 §1：推文总数 ~955,012 条，覆盖 HongKong、Iran；§7.2：npm_data.json 5,287 节点/4,191 边。'],
    ],
)

doc.save(OUT)
print('Done! Output: ' + OUT)
