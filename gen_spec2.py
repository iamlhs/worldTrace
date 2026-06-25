# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'd:\华东师大计科\6月事件文件\新系统\file_source\EDAS事件分析模块_软件开发说明书.docx'

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0); section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.name='黑体'; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0x1a,0x1a,0x1a)
def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name='黑体'; r.font.size=Pt(15); r.font.color.rgb=RGBColor(0x1a,0x1a,0x1a)
def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name='黑体'; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0x1a,0x1a,0x1a)

def p(text, indent=True, bold=''):
    para = doc.add_paragraph()
    if indent: para.paragraph_format.first_line_indent = Cm(0.74)
    if bold:
        r=para.add_run(bold); r.bold=True; r.font.name='宋体'; r.font.size=Pt(10.5)
    r=para.add_run(text); r.font.name='宋体'; r.font.size=Pt(10.5)

def bullet(text):
    bp = doc.add_paragraph(style='List Bullet'); bp.clear()
    r=bp.add_run(text); r.font.name='宋体'; r.font.size=Pt(10.5)

def code(text):
    cp = doc.add_paragraph()
    cp.paragraph_format.left_indent = Cm(1.5)
    r=cp.add_run(text); r.font.name='Consolas'; r.font.size=Pt(9.5)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)

# ==================== COVER ====================
for _ in range(6): doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('EDAS 事件分析模块'); r.bold=True; r.font.size=Pt(28); r.font.name='黑体'
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('软件开发说明书'); r.bold=True; r.font.size=Pt(24); r.font.name='黑体'
doc.add_paragraph()
for line in ['版本 V1.0', '2026 年 6 月']:
    l=doc.add_paragraph(); l.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=l.add_run(line); r.font.size=Pt(12); r.font.name='宋体'; r.font.color.rgb=RGBColor(0x66,0x66,0x66)
doc.add_page_break()

# ==================== 1 ====================
h1('1  系统概述')
h2('1.1  项目背景')
p('EDAS（Event Detection and Analysis from Social Media）是基于社交媒体数据的开源事件检测与分析系统。该系统从海量推文中自动检测热点事件、生成摘要、提取关键实体并构建关系图谱。本事件分析模块是该系统的核心前端分析引擎，聚焦于三个关键技术方向：基于知识图谱的事件目标组织关联关系分析、基于地理空间编码的事件空间维度可视化、基于深度自然语言处理的事件目标多维度研判。')
h2('1.2  系统定位')
p('事件分析模块面向地缘政治情报分析场景，覆盖香港、伊朗、乌克兰三大热点区域的事件数据。系统以深度学习与知识图谱技术为底座，提供从事件自动检测、实体关系挖掘到多维信息关联展现的完整分析链路，支撑情报分析人员的快速态势感知与深度推理需求。')
h2('1.3  核心数据规模')
# Data table
dt = doc.add_table(rows=4, cols=2); dt.style='Table Grid'
for row,w in zip(dt.rows, [['数据指标','数量'],['推文总数','约 955,012 条'],['覆盖地区','香港、伊朗、乌克兰'],['检测事件','474 个']]):
    for j,val in enumerate(w):
        c=row.cells[j]; c.text=''; r=c.paragraphs[0].add_run(val)
        r.font.name='宋体'; r.font.size=Pt(10)
        if row==dt.rows[0]: r.bold=True
doc.add_paragraph()
doc.add_page_break()

# ==================== 2 ====================
h1('2  核心技术栈')
p('事件分析模块基于成熟的深度学习与自然语言处理技术栈构建，以下分别阐述各关键技术在本系统中的实际应用。')
h2('2.1  事件摘要生成：BART 预训练模型')
p('BART（Bidirectional and Auto-Regressive Transformer）是 Meta 提出的序列到序列预训练模型，融合了 BERT 的双向编码器与 GPT 的自回归解码器。在本系统中，BART 用于事件的抽象式摘要生成（Abstractive Summarization），属于离线预处理阶段的核心算法。')
p('',bold='技术原理：')
bullet('编码器采用双向自注意力机制（Bi-directional Self-Attention），对输入的推文集合进行全局上下文编码，捕捉跨句子的语义关联。')
bullet('解码器采用自回归生成策略，通过 Masked Self-Attention 逐词生成摘要文本，同时利用 Cross-Attention 机制关注编码器的隐层表示，确保生成内容忠实于原文信息。')
bullet('与抽取式摘要不同，BART 能够重新组织语言、融合多个子事件的信息，生成语法正确、语义连贯的自然语言摘要。')
p('在系统中的应用流程：推文采集 → 文本清洗与拼接（截断至 1024 tokens）→ BART 编码 → 自回归解码生成摘要 → 存储为事件描述字段。每个事件均携带由 BART 生成的摘要文本，作为后续事件分类、关键词提取和因果推理的基础输入。')

h2('2.2  命名实体识别：BiLSTM-CRF 序列标注模型')
p('系统采用 PyTorch 框架构建的 BiLSTM-CRF 深度序列标注模型，用于从推文文本中自动抽取人物（Person）、地点（Location）、组织（Organization）等命名实体。')
p('',bold='模型架构：')
bullet('词嵌入层：将输入文本中的每个词映射为稠密向量表示，捕获词的语义信息。')
bullet('BiLSTM 编码层：双向长短期记忆网络分别从前向和后向扫描文本序列，捕捉每个词的上下文依赖关系，生成包含全局上下文信息的隐状态表示。')
bullet('CRF（条件随机场）解码层：在 BiLSTM 输出的发射概率基础上，学习标签之间的转移约束（例如 I-Person 不能跟在 B-Location 之后），通过维特比算法求解全局最优标签序列。')
p('抽取的命名实体直接服务于知识图谱的节点构建——人物实体成为图谱中的人物节点，地点实体映射到空间维度的地理位置信息，组织实体参与事件目标组织关联关系的建立。')

h2('2.3  事件类型分类与突发检测')
p('事件类型的自动识别是目标维度分析的基础。系统综合运用多种技术手段实现事件的分类与突发性判定。')
p('',bold='TensorFlow 事件分类模型：')
p('系统集成了基于 TensorFlow 1.x 框架的预训练文本分类模型。模型采用静态图（Static Graph）模式加载推理，支持对推文文本进行事件类别判定。典型架构为 Embedding → Conv1D/LSTM → Dense → Softmax 的多层神经网络，输出事件属于各类别的概率分布。')
p('',bold='峰值检测算法（Peak Detection）：')
p('系统内置了基于信号处理的时序峰值检测算法，用于识别推文数量时间序列中的突发峰值。算法通过扫描一维时间序列数据，在设定的峰值最小间距（Peak Distance）约束下检测局部最大值，从而判定当前时刻是否发生突发性事件。峰值检测结果直接映射为事件的突发标志（Bursty Flag），为情报分析人员提供事件紧急程度的量化参考。')

doc.add_page_break()

# ==================== 3 ====================
h1('3  事件目标组织关联关系图')
h2('3.1  功能概述')
p('事件目标组织关联关系图（知识图谱）是事件分析模块的重要组成部分。系统基于从推文事件中抽取的实体及实体间关系，构建多类型的关系图谱，以力导向图的方式直观展示事件中涉及的人物、组织、装备等目标实体之间的关联网络。')
h2('3.2  知识图谱的数据构建')
p('知识图谱数据来源于离线预处理阶段对推文文本的深度挖掘。通过命名实体识别（BiLSTM-CRF）从推文中抽取人物、地点、组织等实体，结合实体共现分析和关系抽取算法，构建实体之间的关联边。图谱数据以 JSON 和 GEXF 格式持久化存储，并通过 MongoDB 数据库进行管理和查询。')
p('系统预置三条核心知识图谱链：')
bullet('政治人物关联链（Trump→Pence）：包含 103 个节点，展示美国政治人物之间的 Vice President、Successor、Leader、Deputy 等关系类型。')
bullet('政治人物关联链（Obama→Biden）：包含 114 个节点，覆盖更广泛的政治人物关系网络。')
bullet('卫星任务链（NOAA）：包含 44 个节点，展示 NOAA 卫星系列中 Previous Mission 和 Next Mission 的任务承接关系。')
p('每条关系边均附带关系类型标签，支持用户直观理解目标实体之间的语义关联。')
h2('3.3  图谱的物理仿真渲染')
p('知识图谱的交互式可视化基于力导向图（Force-Directed Graph）物理仿真模型。该模型将图谱中的节点视为带电粒子、边视为弹簧连接，通过迭代求解力学平衡方程实现节点的自动布局：')
bullet('连接力（Link Force）：通过胡克定律模拟，将有关联的节点吸引至理想距离，确保关系紧密的节点在空间中彼此靠近。')
bullet('排斥力（Many-Body Force）：基于库仑定律模拟，所有节点之间施加电荷排斥力，避免节点过度重叠，保证标签可读性。')
bullet('向心力（Center Force）：将所有节点向画布中心吸引，防止力模型发散导致节点漂移至视野之外。')
bullet('碰撞力（Collision Force）：在节点间施加碰撞半径约束，确保最小间距。')
p('力模型通过多次迭代（Tick）逐步收敛至能量最小化的稳定状态。首次收敛完成时，系统记录并展示全流程渲染耗时，用于验证图谱生成效率。')

doc.add_page_break()

# ==================== 4 ====================
h1('4  事件空间维度分析')
h2('4.1  功能概述')
p('事件空间维度分析以电子地图为载体，将 EDAS 系统检测到的事件按地理坐标精确映射到全球地图上。系统覆盖香港、伊朗、乌克兰三大地缘热点区域，每个区域的事件以独立颜色体系编码，支持多尺度空间聚类与信息钻取。')
h2('4.2  空间数据编码')
p('每个事件均携带由地理位置处理模块（Location Processing）提取的经纬度坐标和地区归属信息。坐标数据来源于推文的 Geo-tag 元数据或文本中地名的地理编码（Geocoding）。系统为三大区域分别建立空间索引，实现按区域的事件快速检索与统计。')
h2('4.3  多尺度空间聚类')
p('为解决大规模事件点在地图上的视觉重叠问题，系统采用基于空间索引的层次聚类算法。该算法在低缩放级别将空间邻近的事件点聚合成聚类圆，聚类圆的半径随内部事件数量动态增长（对数尺度），颜色由聚类内主导区域决定，并以文字标注聚类内事件总数。当用户放大地图至高级别时，聚类圆自动拆分为独立事件点，实现从宏观概览到微观细节的无缝过渡。')
h2('4.4  空间维度信息关联')
p('每个事件的空间维度信息包含经度、纬度、区域归属和具体位置名称。用户通过点击地图上的事件点或聚类圆，可查看精确的地理坐标信息（经纬度保留至小数点后四位）和所在地名称。这些空间信息与事件的目标维度信息（事件类型、等级、关键词）形成交叉关联，支撑多维度的情报综合分析。')

doc.add_page_break()

# ==================== 5 ====================
h1('5  事件目标维度分析')
h2('5.1  功能概述')
p('事件目标维度分析是系统的核心分析引擎，为每个事件提供多层次、多维度的目标信息研判决。系统综合运用文本分类、关键词提取、因果推理等技术，自动生成事件的多维分析结果。')
h2('5.2  事件类型自动分类')
p('系统基于多规则推理引擎实现四类事件类型的自动识别：抗议类、冲突类、军事活动类和其他类。分类引擎融合了两类特征：其一，基于关键词密度的语义特征——每类事件预设核心关键词库（25-35 个），计算事件文本与各类关键词库的匹配密度；其二，基于区域分布的先验特征——不同地缘区域的事件类型分布存在先验差异，系统以温和的偏置权重（最大 +0.5）引入区域信息。综合两类特征判定最终事件类型，确保分类结果兼具语义准确性和场景适应性。')
h2('5.3  事件等级与突发性判定')
p('事件等级划分为特别重大事件、重大事件、较大事件和一般事件四个层级。等级判定综合了事件关键词密度、峰值检测算法输出的突发概率、以及事件在时间序列中的相对显著性。突发标志直接由峰值检测算法的输出映射——当推文时间序列在事件时刻检测到显著峰值时，该事件被标记为"突发"，为情报分析人员提供紧急程度预警。')
h2('5.4  关键词权重分析与语义聚焦')
p('系统对每个事件的摘要文本执行基于注意力机制的关键词权重分析。通过 BART 模型编码器中的自注意力权重矩阵，计算每个 token 对事件语义表示的贡献度，提取 Top-12 高权重关键词。分析结果以权重条形图和词云两种形式呈现：条形图直观对比各关键词的相对重要性，词云以字体大小和颜色深浅编码权重高低，揭示模型关注的核心语义焦点。')
h2('5.5  事件因果链推理')
p('因果链推理是目标维度分析的高级功能，旨在自动发现同一区域内事件之间的因果关联。系统采用基于共享关键词的有向图推理算法：')
bullet('共享关键词筛选：两个事件至少共享 4 个高权重关键词，确保语义层面的关联性。')
bullet('双向重叠度约束：共享关键词在两个事件各自的关键词集合中的占比均不低于 30%，避免单向弱关联。')
bullet('时间窗口约束：两事件的时间间隔不超过 60 天，确保因果关联的时序合理性。')
bullet('边数控制：每个方向最多保留 6 条最强关联边，避免过度连接导致的图结构混乱。')
p('因果链以层级可折叠结构呈现，每条关联边标注共享关键词标签、双向重叠率和时间间隔，支持用户逐层展开探索事件间的因果传导关系。')

doc.add_page_break()

# ==================== 6 ====================
h1('6  数据存储与系统集成')
h2('6.1  知识图谱数据存储')
p('知识图谱的实体和关系数据采用 MongoDB 文档数据库进行持久化存储。MongoDB 的文档模型天然适合存储图数据的节点-边结构，支持灵活的 Schema 设计和高效的嵌套查询。图谱数据同时以 JSON 和 GEXF（Graph Exchange XML Format）格式保存静态副本，支持离线分析和第三方工具导入。')
h2('6.2  事件数据管理')
p('事件元数据（摘要、关键词、分类结果、时空属性）以 JSON 格式组织，存储于文件系统中。每个事件包含完整的分析字段：事件 ID、日期、经纬度坐标、区域归属、位置名称、事件等级、突发标志、类型分类、关键词及权重、摘要文本等。前端通过 HTTP API 按需加载事件数据，支持区域过滤和分页查询。')
h2('6.3  与平台其他模块的关系')
p('事件分析模块作为智能情报分析平台的核心组件，与平台的其他模块形成数据联动：视频理解模块提供的事件目标研判结果可作为事件目标维度的补充信息；金融舆情模块的事件演化趋势预测模型可为本模块的因果链推理提供时序先验知识；医疗事件模块的可解释性技术（注意力机制、元路径、图神经网络、认知机制）可为事件分析模块的方法论提供参考。')

doc.add_page_break()

# ==================== 7 ====================
h1('7  部署与运行')
h2('7.1  软件环境')
bullet('Python ≥ 3.10，含 PyTorch、Transformers、TensorFlow 等深度学习框架')
bullet('MongoDB ≥ 4.0（知识图谱数据存储）')
bullet('Node.js ≥ 18.x（前端开发服务器）')
bullet('现代浏览器（支持 WebGL 1.0+）')
h2('7.2  启动方式')
code('npm run dev          # 启动 Vite 开发服务器')
code('python server.py    # 启动 Python 数据 API 服务')
h2('7.3  数据准备')
p('事件分析模块依赖预处理的 EDAS 事件数据集，包括事件 JSON 文件和知识图谱 JSON/GEXF 文件。数据集存放于 edas_exports 目录下，系统启动时自动加载。图谱数据需预先导入 MongoDB 数据库 cbd109KG 中。')

doc.save(OUT)
print('Done! Output: ' + OUT)
