# -*- coding: utf-8 -*-
"""解析Word文档表格格式"""
import docx

doc = docx.Document(r'd:\华东师大计科\6月事件文件\新系统\功能测试表例子.docx')

print('=== 段落样式 ===')
for i, p in enumerate(doc.paragraphs[:10]):
    runs = p.runs
    font_name = runs[0].font.name if runs else 'N/A'
    font_size = runs[0].font.size if runs else 'N/A'
    bold = runs[0].font.bold if runs else 'N/A'
    print(f'P[{i}]: text={p.text[:80]}')
    print(f'       style={p.style.name} | align={p.alignment} | font={font_name} | size={font_size} | bold={bold}')

print()
print(f'=== 表格数量: {len(doc.tables)} ===')

for ti, table in enumerate(doc.tables):
    print(f'\n--- Table {ti} ---')
    print(f'  Rows: {len(table.rows)}, Cols: {len(table.columns)}')

    # Column widths
    widths = []
    for ci, col in enumerate(table.columns):
        w = col.width
        widths.append(f'Col{ci}={w}')
    print(f'  Col widths: {", ".join(widths)}')

    # Print all rows structure
    for ri, row in enumerate(table.rows):
        cells_info = []
        for ci, cell in enumerate(row.cells):
            txt = cell.text.replace('\n', ' ')[:50]
            tc = cell._tc
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            tcPr = tc.find(f'{ns}tcPr')
            span_val = None
            if tcPr is not None:
                grid_span = tcPr.find(f'{ns}gridSpan')
                if grid_span is not None:
                    span_val = grid_span.get(f'{ns}val')

            # Check cell shading
            shading = None
            if tcPr is not None:
                shd = tcPr.find(f'{ns}shd')
                if shd is not None:
                    shading = shd.get(f'{ns}fill')

            # Check paragraph alignment in cell
            para = cell.paragraphs[0] if cell.paragraphs else None
            align = para.alignment if para else None

            # Check font properties
            font_info = ''
            if para and para.runs:
                r = para.runs[0]
                font_info = f'bold={r.font.bold},size={r.font.size},name={r.font.name}'

            cells_info.append(f'C{ci}[span={span_val},shd={shading},align={align},{font_info}]:{txt}')
        print(f'  Row[{ri}]: {" | ".join(cells_info)}')
