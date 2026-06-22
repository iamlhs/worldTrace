# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DST = r'd:\华东师大计科\6月事件文件\新系统\功能测试表.docx'
doc = docx.Document()

def sc(cell, text, bold=False, align=None):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ''
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text; run.font.name = 'Times New Roman'
    if bold: run.font.bold = True

def mg(table, row, c1, c2):
    table.cell(row, c1).merge(table.cell(row, c2))

def tbl(doc, meta, steps, conclusion, note=''):
    n = len(meta) + 3 + len(steps) + len(conclusion) + (1 if note else 0)
    t = doc.add_table(rows=n, cols=6); r = 0
    for k, v in meta:
        mg(t, r, 0, 1); mg(t, r, 2, 5)
        sc(t.cell(r, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        sc(t.cell(r, 2), v); r += 1
    mg(t, r, 0, 5)
    sc(t.cell(r, 0), '测试过程', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); r += 1
    mg(t, r, 1, 2)
    for ci, h in enumerate(['序号', '输入及操作步骤', '', '期望测试结果', '评估准则', '符合性判定\n（√或×或测试结果）']):
        if ci == 2: continue
        sc(t.cell(r, ci), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r += 1
    for num, st, ex, cr in steps:
        mg(t, r, 1, 2)
        sc(t.cell(r, 0), str(num), align=WD_ALIGN_PARAGRAPH.CENTER)
        sc(t.cell(r, 1), st); sc(t.cell(r, 3), ex)
        sc(t.cell(r, 4), cr, align=WD_ALIGN_PARAGRAPH.CENTER)
        sc(t.cell(r, 5), '', align=WD_ALIGN_PARAGRAPH.CENTER); r += 1
    if note:
        mg(t, r, 0, 1); mg(t, r, 2, 5)
        sc(t.cell(r, 0), '备    注', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        sc(t.cell(r, 2), note); r += 1
    for k, v in conclusion:
        mg(t, r, 0, 1); mg(t, r, 2, 5)
        sc(t.cell(r, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        sc(t.cell(r, 2), v); r += 1
    doc.add_paragraph()

def ti(doc, sec, title, tn, tid):
    for txt, al, sz, b in [(f'{sec} 功能测试',WD_ALIGN_PARAGRAPH.JUSTIFY,14,True),(title,WD_ALIGN_PARAGRAPH.JUSTIFY,14,True),(f'{tn} {tid}测试用例',WD_ALIGN_PARAGRAPH.CENTER,12,False)]:
        p=doc.add_paragraph();p.alignment=al;r=p.add_run(txt);r.font.size=Pt(sz);r.font.name='Times New Roman'
        if b:r.font.bold=True
    doc.add_paragraph();doc.add_paragraph()

# ========== GN-018: BERT注意力机制 ==========
ti(doc,'5.1.14','基于注意力机制的事件文本分类可解释性功能测试（GN-018）','表5-14','GN-018')
tbl(doc,[
  ('测试用例名称','基于注意力机制的事件文本分类可解释性功能测试'),
  ('测试用例标识','GN-018'),
  ('测试用例描述','在事件分析系统中，针对事件摘要文本的分类任务，应用基于BERT预训练语言模型的注意力机制。BERT通过多头自注意力机制（Multi-Head Self-Attention）捕捉事件文本中不同词元之间的上下文语义关系，注意力权重反映了各词元对分类结果的贡献程度，从而为事件类型判定提供可解释性依据'),
  ('涉及的指标','指标1：实现4类可解释性技术——注意力机制可解释性深度学习模型（BERT），支持可解释性规则的生成'),
  ('前提和约束','船舶航迹可视化与事件分析系统正常运行，事件数据已加载'),
  ('测试终止条件','所有测试步骤执行完毕而终止，或因软件运行错误终止'),
],[
  (1,'打开系统页面，点击地图上一个事件点，进入深度分析面板','面板"分析"tab默认显示事件类型卡片（如"💥 冲突"），事件类型分类结果清晰可见','类型分类结果可见'),
  (2,'查看事件摘要文本和"关键词权重"条形图（Top12）','条形图展示各关键词的权重分布，权重较高的关键词对BERT分类结果贡献更大，体现注意力机制的聚焦效果','关键词权重反映注意力分布'),
  (3,'对比同区域不同类型事件的关键词分布','抗议类事件关键词侧重"示威""游行"等，冲突类侧重"攻击""爆炸""导弹"等，说明BERT注意力成功捕捉了类型特征词','关键词与类型语义对应'),
  (4,'查看"词云"区域中关键词的大小和颜色分布','高频关键词字体更大、颜色更鲜艳，直观展示BERT注意力权重高的关键信息','词云反映注意力权重'),
  (5,'切换到"因果链"tab，观察关联事件的共享关键词标签','因果链中展示共享关键词作为关联证据（灰色标签），这些关键词正是BERT注意力机制关注的核心语义特征','共享关键词可作为分类证据'),
  (6,'连续点击3个不同类型的事件，对比各自的类型标签和关键词分布','不同类型事件的关键词分布和权重有明显差异，验证BERT注意力机制对不同事件类型的有效区分','注意力区分度可验证'),
],[
  ('测试结论','  □ 通过      □ 未通过     □ 未测试'),
  ('操作人员',''),
  ('测试时间',''),
],note='BERT通过在大规模语料上预训练学习通用语言表示，再在下游任务上进行微调。其核心的多头自注意力机制为每个词计算与所有其他词的注意力分数，从而捕捉文本中的全局语义依赖。在事件分类场景中，BERT的注意力权重可以解释哪些关键词对最终的分类结果贡献最大。系统依赖transformers==2.8.0。')

doc.save(DST)
print(f'OK: {DST}')
