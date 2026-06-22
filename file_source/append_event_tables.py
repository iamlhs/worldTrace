# -*- coding: utf-8 -*-
"""在功能测试表例子.docx末尾追加3张EDAS事件研判测试表"""
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'd:\华东师大计科\6月事件文件\新系统\功能测试表例子.docx'
DST = r'd:\华东师大计科\6月事件文件\新系统\功能测试表例子_事件研判.docx'

doc = docx.Document(SRC)

def set_cell(cell, text, bold=False, align=None, size=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    run.font.name = 'Times New Roman'
    if bold:
        run.font.bold = True
    if size:
        run.font.size = Pt(size) if isinstance(size, (int, float)) else size

def merge(table, row, c1, c2):
    table.cell(row, c1).merge(table.cell(row, c2))

def build_table(doc, meta, steps, conclusion):
    rows = len(meta) + 2 + len(steps) + len(conclusion)
    t = doc.add_table(rows=rows, cols=6)
    r = 0
    # Meta (key spans 0-1, value spans 2-5)
    for k, v in meta:
        merge(t, r, 0, 1)
        merge(t, r, 2, 5)
        set_cell(t.cell(r, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 2), v)
        r += 1
    # 测试过程
    merge(t, r, 0, 5)
    set_cell(t.cell(r, 0), '测试过程', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r += 1
    # Header
    merge(t, r, 1, 2)
    hdrs = ['序号', '输入及操作步骤', '', '期望测试结果', '评估准则', '符合性判定\n（√或×或测试结果）']
    for ci, h in enumerate(hdrs):
        if ci == 2: continue
        set_cell(t.cell(r, ci), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    r += 1
    # Steps
    for num, step_text, expect, criteria, check in steps:
        merge(t, r, 1, 2)
        set_cell(t.cell(r, 0), str(num), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 1), step_text)
        set_cell(t.cell(r, 3), expect)
        set_cell(t.cell(r, 4), criteria, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 5), check, align=WD_ALIGN_PARAGRAPH.CENTER)
        r += 1
    # Conclusion
    for k, v in conclusion:
        merge(t, r, 0, 1)
        merge(t, r, 2, 5)
        set_cell(t.cell(r, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 2), v)
        r += 1
    doc.add_paragraph()
    return t

def add_title(doc, section, title, table_num, test_id):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f'{section} 功能测试')
    r.font.size = Pt(14); r.font.bold = True; r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(title)
    r.font.size = Pt(14); r.font.bold = True; r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{table_num} {test_id}测试用例')
    r.font.size = Pt(12); r.font.name = 'Times New Roman'

    doc.add_paragraph(); doc.add_paragraph()

# ======== Table 1: GN-015 ========
add_title(doc, '5.1.11', '基于时空事件数据的事件目标属性研判功能测试（GN-015）', '表5-11', 'GN-015')
build_table(doc, [
    ('测试用例名称', '基于时空事件数据的事件目标属性研判功能测试'),
    ('测试用例标识', 'GN-015'),
    ('测试用例描述', '本测试用例验证系统对时空事件数据的属性研判能力，包括事件类型分类（抗议/冲突/军事活动/其他）、事件等级（特别重大/重大/较大/一般）、突发标记识别、地理位置标注以及日期时序分析。验收人通过点击事件点进入分析面板，逐项确认事件属性信息的完整性和准确性'),
    ('涉及的指标', '指标4研判任务：事件目标属性研判——展示事件类型、等级、突发标记、位置、日期等属性'),
    ('前提和约束', '船舶航迹可视化与事件分析系统正常运行，事件数据已加载（地图上可见事件聚类圆/事件点）'),
    ('测试终止条件', '所有测试步骤执行完毕而终止，或因软件运行错误终止'),
], [
    (1, '打开系统页面，确认地图上可见事件聚类圆（带数字的彩色圆圈）', '聚类圆可见，数字代表该位置聚合的事件数量，三个区域（香港/伊朗/乌克兰）均有分布', '事件数据加载成功', ''),
    (2, '点击一个事件聚类圆，在弹出的列表中观察每个事件行的信息', '每行显示区域标识、事件类型标签（抗议/冲突/军事活动/其他）、日期、突发标记、等级标签', '属性标签齐全', ''),
    (3, '点击列表中第一条事件，进入深度分析面板', '面板顶部显示醒目的类型卡片（含图标+类型名称），下方展示等级标签、突发标记、位置信息', '属性展示完整', ''),
    (4, '观察分析面板中的事件摘要和关键词', '摘要完整显示事件描述，关键词以条形图和词云形式展示权重分布', '摘要+关键词可见', ''),
    (5, '按ESC关闭面板，点击不同区域的事件聚类圆（如从香港切换到乌克兰）', '不同区域事件属性有差异（类型、关键词分布不同），面板正确更新', '区域切换正常', ''),
    (6, '分别记录3个不同事件的事件类型和等级', '至少2个事件的事件类型不同（不全是同一类型），等级标签与事件严重程度匹配', '类型有区分度', ''),
], [
    ('测试结论', '  □ 通过      □ 未通过     □ 未测试'),
    ('操作人员', ''),
    ('测试时间', ''),
])

# ======== Table 2: GN-016 ========
add_title(doc, '5.1.12', '基于时空事件数据的事件目标行为研判功能测试（GN-016）', '表5-12', 'GN-016')
build_table(doc, [
    ('测试用例名称', '基于时空事件数据的事件目标行为研判功能测试'),
    ('测试用例标识', 'GN-016'),
    ('测试用例描述', '本测试用例验证系统对时空事件行为模式的研判能力，包括事件的时间因果链分析（事件间的有向因果关系）、事件聚类空间分布、同区域事件的时序关联。验收人通过因果链tab和聚类交互，验证系统对事件行为模式的识别和展示能力'),
    ('涉及的指标', '指标4研判任务：事件目标行为研判——通过因果链展示事件间的有向因果关系（前置事件→当前事件→结果事件），支持1对1/1对多/多对1/多对多模式'),
    ('前提和约束', '船舶航迹可视化与事件分析系统正常运行，事件数据已加载'),
    ('测试终止条件', '所有测试步骤执行完毕而终止，或因软件运行错误终止'),
], [
    (1, '点击事件点打开深度分析面板，切换到"因果链"tab', '面板切换为因果链视图，显示"引发当前事件的前置事件"折叠按钮和"当前事件可能引发的结果事件"折叠按钮', '因果链tab正常', ''),
    (2, '点击"引发当前事件的前置事件"按钮展开列表', '展示与当前事件有强因果关联的前置事件列表，每个事件显示日期、类型标签、摘要和共享关键词', '前置事件有向展示', ''),
    (3, '点击"当前事件可能引发的结果事件"按钮展开列表', '展示当前事件可能引发的后继事件列表，每个事件同样显示日期、类型标签、摘要和共享关键词', '结果事件有向展示', ''),
    (4, '观察因果链中事件间的箭头方向（前置→当前/当前→结果）', '箭头明确指示因果方向：前置事件指向当前，当前事件指向结果，非双向混淆', '方向明确', ''),
    (5, '选择关键词较多的热点事件查看因果链', '热点事件的因果关联数量多于冷门事件，体现了事件行为的传播特性', '关联数量合理', ''),
    (6, '选择关键词较少的事件查看因果链', '显示"该事件暂无检测到的因果关联"，不崩溃不报错', '空数据处理正常', ''),
], [
    ('测试结论', '  □ 通过      □ 未通过     □ 未测试'),
    ('操作人员', ''),
    ('测试时间', ''),
])

# ======== Table 3: GN-017 ========
add_title(doc, '5.1.13', '基于时空事件数据的事件目标意图识别研判功能测试（GN-017）', '表5-13', 'GN-017')
build_table(doc, [
    ('测试用例名称', '基于时空事件数据的事件目标意图识别研判功能测试'),
    ('测试用例标识', 'GN-017'),
    ('测试用例描述', '本测试用例验证系统对事件意图和演化趋势的识别研判能力。通过因果链的有向关系推断事件的可能演化方向，结合事件类型分类和等级标签判断事件的发展态势，帮助分析人员识别事件的潜在意图和后续发展趋势'),
    ('涉及的指标', '指标4研判任务：事件目标意图识别研判——通过因果链有向关系+事件类型+等级综合研判事件意图和演化趋势'),
    ('前提和约束', '船舶航迹可视化与事件分析系统正常运行，事件数据已加载'),
    ('测试终止条件', '所有测试步骤执行完毕而终止，或因软件运行错误终止'),
], [
    (1, '打开深度分析面板，查看"分析"tab中的事件类型、等级和摘要', '类型卡片醒目展示（如"冲突"或"抗议"），等级标签和突发标记并列显示', '事件类型和等级明确', ''),
    (2, '切换到"因果链"tab，观察前置事件和结果事件的数量对比', '如果结果事件数量多于前置事件数量，说明该事件可能是"引爆点"；反之可能是"结果事件"', '因果方向可研判意图', ''),
    (3, '观察因果链中前置事件的类型分布', '如果前置事件以"军事活动"为主、当前事件为"冲突"，可研判为军事部署导致冲突升级的意图链', '类型链可辅助研判', ''),
    (4, '观察因果链中结果事件的类型分布', '如果结果事件中包含多个"冲突"类事件，可研判为当前事件可能引发连锁冲突的演化趋势', '演化趋势可预测', ''),
    (5, '对比同一区域内不同事件在因果链中的位置（前置/当前/结果）', '不同事件在因果网络中扮演不同角色（起因/转折/结果），体现事件意图的多样性', '角色定位有区分', ''),
    (6, '综合事件类型+等级+因果链方向，对某一事件写出研判结论', '验收人可根据系统提供的信息（类型+等级+因果方向+共享关键词）得出合理的事件意图研判', '综合研判可行', ''),
], [
    ('测试结论', '  □ 通过      □ 未通过     □ 未测试'),
    ('操作人员', ''),
    ('测试时间', ''),
])

doc.save(DST)
print(f'Saved: {DST}')
