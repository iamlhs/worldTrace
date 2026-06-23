import docx

for name, label in [
    ('AI软件需收集模型及调研信息_金融.docx', '金融'),
    ('AI软件需收集模型及调研信息_中医.docx', '中医'),
    ('AI软件需收集模型及调研信息_视频.docx', '视频'),
]:
    path = rf'd:\华东师大计科\6月事件文件\新系统\{name}'
    doc = docx.Document(path)
    print(f'\n===== {label} =====')
    # Find the last table (功能对照表)
    if doc.tables:
        t = doc.tables[-1]
        print(f'Last table: {len(t.rows)}r x {len(t.columns)}c')
        for ri, row in enumerate(t.rows):
            cells = [c.text.replace('\n',' ')[:60] for c in row.cells]
            print(f'  R{ri}: ' + ' | '.join(cells))
    # Also print section near 功能对照表
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if '功能对照' in t or '功能清单' in t:
            print(f'P[{i}]: {t[:100]}')
