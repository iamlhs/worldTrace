import docx
doc = docx.Document(r'd:\华东师大计科\6月事件文件\新系统\AI软件需收集模型及调研信息.docx')
print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'P[{i}]: {p.text[:150]}')
print()
print(f'=== TABLES: {len(doc.tables)} ===')
for ti, table in enumerate(doc.tables):
    print(f'\n--- Table {ti} ({len(table.rows)}r x {len(table.columns)}c) ---')
    for ri, row in enumerate(table.rows):
        cells = []
        for c in row.cells:
            txt = c.text.replace('\n', ' ')[:80]
            cells.append(txt)
        print(f'  R{ri}: ' + ' | '.join(cells))
