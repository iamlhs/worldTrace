import docx

# Read 软测文档
doc = docx.Document(r'd:\华东师大计科\6月事件文件\新系统\软测需提供材料和信息.docx')
print('=== 软测文档段落 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f'P[{i}]: {t[:200]}')

print(f'\n=== 软测文档表格 ({len(doc.tables)} tables) ===')
for ti, t in enumerate(doc.tables):
    print(f'T{ti}: {len(t.rows)}r x {len(t.columns)}c')
    for ri, row in enumerate(t.rows):
        cells = []
        for c in row.cells:
            cells.append(c.text.replace('\n', ' ')[:80])
        print(f'  R{ri}: ' + ' | '.join(cells))
    print()
