# -*- coding: utf-8 -*-
"""修改AI软件需收集模型及调研信息.docx — 补充船舶航迹系统模型和数据集"""
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

SRC = r'd:\华东师大计科\6月事件文件\新系统\AI软件需收集模型及调研信息.docx'
DST = r'd:\华东师大计科\6月事件文件\新系统\AI软件需收集模型及调研信息_已填充.docx'

doc = docx.Document(SRC)

# ============================================================
# Helper: find paragraph by text prefix and replace its text
# ============================================================
def find_and_replace(prefix, new_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            # Clear all runs and set new text
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            return p
    print(f'WARNING: not found: {prefix[:30]}...')
    return None

# ============================================================
# 1. 替换模型描述
# ============================================================
find_and_replace('本系统包括文本特征提取算法3种',
    '本系统包括船舶轨迹预测深度学习模型4种，分别为门控循环单元(GRU)、循环神经网络(RNN)、长短期记忆网络(LSTM)和时空点过程模型(DSTPP)，用于对船舶航行轨迹进行建模与预测，并结合EDAS事件分析实现态势研判。')

# GRU
find_and_replace('1）HM-DenseRNNs',
    '1）GRU（门控循环单元，Gated Recurrent Unit）')
find_and_replace('HM-DenseRNNs包括密集连接循环神经网络',
    'GRU是一种通过更新门(Update Gate)和重置门(Reset Gate)控制信息流动的循环神经网络变体。相比LSTM参数更少、训练更快，在船舶轨迹预测中用于捕捉航迹点之间的时序依赖关系，输出下一时刻的位置预测。模型接收历史航迹点序列（经纬度+时间），通过门控机制选择性保留或遗忘历史信息，最终输出未来轨迹的空间位置。')

# RNN
for p in doc.paragraphs:
    if p.text.strip().startswith('2）基于记忆增强神经网络'):
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '2）RNN（循环神经网络，Recurrent Neural Network）'
        break
find_and_replace('基于记忆增强神经网络包括文本初始化模块',
    'RNN是一种通过隐藏状态递归传递序列信息的神经网络基础架构，擅长捕捉短期时序依赖。在船舶轨迹预测中，RNN将历史航迹点按时间顺序依次输入，每个时刻的隐藏状态融合当前输入与上一步隐藏状态，最终输出完整预测轨迹。其简洁的结构使其在短轨迹场景下具有较高的计算效率。')

# LSTM
for p in doc.paragraphs:
    if p.text.strip().startswith('3）基于键值记忆增强神经网络'):
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '3）LSTM（长短期记忆网络，Long Short-Term Memory）'
        break
find_and_replace('基于键值即以增强神经网络包括卷积神经网络',
    'LSTM是一种通过输入门、遗忘门和输出门三个门控单元控制信息流动的循环神经网络，能够有效捕捉航迹中的长距离时序依赖关系。在船舶轨迹预测中，LSTM对历史航迹点进行选择性记忆和遗忘，保留对远距离预测有用的空间模式，在长轨迹场景下表现出较高的预测精度。')

# Add DSTPP after the LSTM section (insert new paragraphs)
# Find the paragraph containing "图2-3" and add DSTPP after its image
for i, p in enumerate(doc.paragraphs):
    if '图2-3' in p.text:
        # Add DSTPP section
        # We'll add paragraphs by manipulating the XML
        break

# For simplicity, let's replace the image captions and add DSTPP in between
# Actually, let's just append new paragraphs after the last model image
# Find paragraph index of "数据集描述"
ds_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '数据集描述':
        ds_idx = i
        break

if ds_idx:
    # Insert DSTPP paragraphs before 数据集描述
    # Since python-docx doesn't support easy paragraph insertion, 
    # we'll insert via XML manipulation
    body = doc.element.body
    from lxml import etree
    
    def make_paragraph(text, style='Normal'):
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        p_el = etree.SubElement(body, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
        pPr = etree.SubElement(p_el, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        pStyle = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', style)
        r = etree.SubElement(p_el, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        t = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return p_el
    
    # Find the paragraph element for 数据集描述 and insert before it
    target_p = doc.paragraphs[ds_idx]._element
    
    # Insert DSTPP description
    p1 = make_paragraph('4）DSTPP（时空点过程模型，Discrete Spatio-Temporal Point Process）')
    p2 = make_paragraph('DSTPP是一种联合建模空间和时间维度的点过程强度函数的深度学习模型，通过时空耦合的方式预测船舶轨迹。在船舶轨迹预测中，DSTPP同时考虑空间位置分布密度和时间到达强度，通过最大似然估计学习模型参数，在轨迹转弯、变向等复杂场景下相比纯时序模型具有更高的预测精度。')
    p3 = make_paragraph('')
    p4 = make_paragraph('模型架构图如下所示：')
    p5 = make_paragraph('')
    p6 = make_paragraph('图2-4 DSTPP模型架构图')
    p7 = make_paragraph('')
    
    # Insert in reverse order to maintain sequence
    target_p.addprevious(p7)
    target_p.addprevious(p6)
    target_p.addprevious(p5)
    target_p.addprevious(p4)
    target_p.addprevious(p3)
    target_p.addprevious(p2)
    target_p.addprevious(p1)


# ============================================================
# 2. 替换数据集描述
# ============================================================
find_and_replace('适用程序不当数据集',
    '船舶航迹预测数据集')
find_and_replace('模型采用从把手案例网爬取的裁判文书数据',
    '本系统航迹预测模型采用基于AIS（船舶自动识别系统）历史航行数据构建的轨迹数据集，共包含100条独立船舶航迹，每条航迹包含经度、纬度、时间戳等字段，按时间顺序排列构成完整航行轨迹。')
find_and_replace('该数据集包含XXX条司法文书数据，针对"适用程序不当"指标，包含"简易程序"和"复杂程序"2类',
    '该数据集包含100条船舶航迹数据，每条航迹按模型分为4组预测结果（GRU/RNN/LSTM/DSTPP），每组包含实际航迹(Ground Truth)和模型预测航迹(Prediction)，数据集如下表所示：')
find_and_replace('针对本指标对应模型，采用从把手案例网爬取的裁判文书数据',
    '针对航迹预测模型，采用上述100条AIS船舶航迹数据作为数据集。训练过程中每条航迹按时间顺序输入，模型输出后续时刻的空间位置预测。模型采用Adam优化器，loss function为均方误差(MSE)，评估指标包含Total RMSE、Spatial RMSE和Temporal RMSE（共计3个维度）。')

# Intention dataset
find_and_replace('XXXX数据集',
    '船舶航行意图分类数据集')
# Find the second XXXX
found_first = False
for p in doc.paragraphs:
    if p.text.strip().startswith('XXXX数据集'):
        if found_first:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = '船舶航行意图分类数据集'
            break
        found_first = True
    elif p.text.strip().startswith('模型采用从把手案例网'):
        pass  # skip the training description

# Find the paragraph about "案由信息不当" and replace
find_and_replace('该数据集包含XXX条司法文书数据，针对"适用程序不当"指标，包含XXX类法律条文',
    '该数据集包含从船舶航迹AIS数据中提取的意图分析样本，覆盖5类航行意图（侦察监视、运输转场、巡逻预警、其他(4)、其他(5)），每类样本数量充足，用于训练和验证船舶意图识别模型，数据集如下表所示：')

# Find the third placeholder
for p in doc.paragraphs:
    if p.text.strip().startswith('针对本指标对应模型'):
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = '针对意图分类模型，采用上述意图数据集，随机选取每类15条共75条样本作为小样本测试集，其余作为训练集。训练配置：batch size=32，epoch=100，优化器=Adam，loss function=Cross Entropy。'
        break

# EDAS events dataset
find_and_replace('该数据集包含XX条司法文书数据，针对"案由信息不当"指标',
    '该数据集包含从公开来源采集的多源事件数据，覆盖香港、伊朗、乌克兰3个地理区域，经过地理编码和关键词提取后聚合为EDAS事件数据集，数据集如下表所示：')
find_and_replace('对本指标对应模型，采用从把手案例网爬取',
    '针对EDAS事件分析模型，采用上述多源事件数据集，经SuperCluster空间聚类和关键词共现分析构建事件关联网络。知识图谱包含节点≥100个，因果链关联基于关键词重叠度和时间邻近度计算。')


# ============================================================
# 3. 更新表格
# ============================================================
# Table 0: trajectory prediction dataset
t0 = doc.tables[0]
data_t0 = [
    ['类别', '类名', '每类样本数', '训练样本数', '测试样本数'],
    ['轨迹数据', '船舶航迹', '100条', '80条（#0-#79）', '20条（#80-#99）'],
    ['模型', 'GRU/RNN/LSTM/DSTPP', '4种', '—', '—'],
    ['评估指标', 'Total/Spatial/Temporal RMSE', '3维度', '—', '—'],
]
for ri, row_data in enumerate(data_t0):
    for ci, val in enumerate(row_data):
        if ri < len(t0.rows):
            cell = t0.rows[ri].cells[ci]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = val
                else:
                    p.add_run(val)

# Table 1: intention classification dataset  
t1 = doc.tables[1]
data_t1 = [
    ['类别', '类名', '每类样本数'],
    ['0', '🔍 侦察监视', '≥15条'],
    ['1', '🚢 运输转场', '≥15条'],
    ['2', '🛡️ 巡逻预警', '≥15条'],
    ['3', '⚔️ 其他(4)', '≥15条'],
    ['4', '🎯 其他(5)', '≥15条'],
]
for ri, row_data in enumerate(data_t1):
    for ci, val in enumerate(row_data):
        if ri < len(t1.rows) and ci < len(t1.rows[ri].cells):
            cell = t1.rows[ri].cells[ci]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = val
                else:
                    p.add_run(val)

# Table 2: EDAS region distribution
t2 = doc.tables[2]
data_t2 = [
    ['类别', '类名', '每类样本数'],
    ['0', '🇭🇰 香港', 'XX 聚合事件'],
    ['1', '🇮🇷 伊朗', 'XX 聚合事件'],
    ['2', '🇺🇦 乌克兰', 'XX 聚合事件（含推文）'],
    ['3', '知识图谱节点', '≥100 个目标'],
    ['4', '因果链关联', '关键词重叠+时间邻近'],
    ['5', '—', '—'],
    ['6', '—', '—'],
    ['7', '—', '—'],
]
for ri, row_data in enumerate(data_t2):
    for ci, val in enumerate(row_data):
        if ri < len(t2.rows) and ci < len(t2.rows[ri].cells):
            cell = t2.rows[ri].cells[ci]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = val
                else:
                    p.add_run(val)

# ============================================================
# 4. 更新系统介绍
# ============================================================
find_and_replace('XX系统，属于XX类型软件，应用于XXX领域',
    '船舶航迹可视化与EDAS事件分析系统，属于时空数据可视化与情报分析类型软件，应用于船舶行为分析与态势研判领域，实现了船舶航迹预测（GRU/RNN/LSTM/DSTPP 4种模型）、EDAS事件空间聚类与深度分析、知识图谱与因果链可视化、航行意图识别等功能，提供了多维度（时间/空间/目标/标签/分类/关联/事件）的信息关联展现与可解释性分析服务。')

# ============================================================
# 5. 更新功能对照表
# ============================================================
find_and_replace('本次测试的功能测试依据为：',
    '本次测试的功能测试依据为：《船舶航迹可视化与EDAS事件分析原型系统技术指标》（6项技术指标）。')

# ============================================================
# 6. 更新性能测试指标
# ============================================================
find_and_replace('系统支持50个用户并发', '知识图谱渲染耗时 ≤ 1秒（≥100个目标节点）')
find_and_replace('xx业务操作响应时间小于2秒', '因果链计算耗时 ≤ 1秒；航迹切换响应时间 ≤ 500ms')

# ============================================================
# 7. 更新开发工具
# ============================================================
find_and_replace('开发工具：如Eclipse 4.7、Microsoft Visual Studio 2010',
    '开发工具：Visual Studio Code；构建工具：Vite 6.0；语言：TypeScript 5.7')

# ============================================================
# 8. 更新日志
# ============================================================
find_and_replace('有没有操作日志：有/无，记录在数据库中', '系统操作日志通过浏览器Console输出，部署环境可通过Nginx访问日志记录HTTP请求。')

# ============================================================
# 9. 更新开发单位信息
# ============================================================
find_and_replace('开发方全称：xxxx公司', '开发方全称：华东师范大学计算机科学与技术学院')
find_and_replace('地址：上海市xx区xx路xx号', '地址：上海市普陀区中山北路3663号')

# ============================================================
# Save
# ============================================================
doc.save(DST)
print(f'Saved to: {DST}')
