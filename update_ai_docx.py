# -*- coding: utf-8 -*-
"""在AI软件需收集模型及调研信息_已填充.docx中补充EDAS事件DL技术描述"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

SRC = r'd:\华东师大计科\6月事件文件\新系统\file_source\AI软件需收集模型及调研信息_已填充.docx'
DST = r'd:\华东师大计科\6月事件文件\新系统\AI软件需收集模型及调研信息_已填充_含事件DL.docx'

doc = docx.Document(SRC)
body = doc.element.body
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def mk_p(text, bold=False, size=Pt(10.5)):
    p_el = etree.SubElement(body, f'{{{W}}}p')
    pPr = etree.SubElement(p_el, f'{{{W}}}pPr')
    if bold:
        b = etree.SubElement(pPr, f'{{{W}}}b')
    r_el = etree.SubElement(p_el, f'{{{W}}}r')
    rPr = etree.SubElement(r_el, f'{{{W}}}rPr')
    sz = etree.SubElement(rPr, f'{{{W}}}sz')
    sz.set(f'{{{W}}}val', str(int(size.pt * 2)))
    t_el = etree.SubElement(r_el, f'{{{W}}}t')
    t_el.text = text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p_el

# Find the paragraph before 数据集描述 to insert EDAS DL section
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '数据集描述':
        target_idx = i
        break

if target_idx:
    target = doc.paragraphs[target_idx]._element
    insertions = [
        ('EDAS事件分析深度学习模型描述', True, Pt(14)),
        ('本系统还包括事件分析深度学习模型4种，用于事件文本的命名实体识别、事件类型分类和语义理解，分别为基于注意力机制的BERT预训练模型、基于图神经网络的事件关联模型、基于元路径的事件知识推理模型和基于认知机制的事件态势研判模型。', False, Pt(10.5)),
        ('1）基于注意力机制的BERT预训练模型', True, Pt(10.5)),
        ('BERT（Bidirectional Encoder Representations from Transformers）通过多头自注意力机制捕捉事件摘要文本中的双向上下文语义，用于事件类型的自动分类（抗议/冲突/军事活动/其他）。注意力权重可解释模型对关键词的关注程度，帮助分析人员理解分类依据。', False, Pt(10.5)),
        ('2）基于图神经网络的事件关联模型', True, Pt(10.5)),
        ('图神经网络通过消息传递机制捕捉事件节点间的图结构关系，学习事件之间的隐含关联模式（关键词共享、时间邻近、空间临近）。在知识图谱中可视化事件节点和关系边，节点重要性排序和边权重提供可解释的关联强度依据。', False, Pt(10.5)),
        ('3）基于元路径的事件知识推理模型', True, Pt(10.5)),
        ('元路径技术定义了实体间的语义关系路径模式（如"事件→关键词→事件"、"事件→位置→事件"、"事件→日期→事件"），通过这些预定义的元路径模式解释事件之间的隐含关联和知识推理链路，为因果链分析提供语义路径级别的可解释性。', False, Pt(10.5)),
        ('4）基于认知机制的事件态势研判模型', True, Pt(10.5)),
        ('认知机制模拟人类分析人员的推理过程，构建"属性认知（事件类型/等级/突发）→ 行为认知（因果链方向）→ 意图演化（趋势预测）"三层递进的研判链路。每一层的分析结果都可追溯到具体的输入数据，实现端到端的可解释态势研判。', False, Pt(10.5)),
        ('', False, Pt(10.5)),
    ]
    # Insert in reverse to maintain order
    for text, bold, size in reversed(insertions):
        target.addprevious(mk_p(text, bold, size))

doc.save(DST)
print(f'Saved: {DST}')
