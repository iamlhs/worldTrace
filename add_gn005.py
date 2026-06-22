# -*- coding: utf-8 -*-
"""在功能测试表例子_事件研判.docx末尾追加 GN-005 事件目标属性研判功能测试"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'd:\华东师大计科\6月事件文件\新系统\功能测试表例子_事件研判.docx'
doc = docx.Document(SRC)

def set_cell(cell, text, bold=False, align=None, size=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    run.font.name = 'Times New Roman'
    if bold: run.font.bold = True
    if size: run.font.size = Pt(size) if isinstance(size, (int, float)) else size

def merge(table, row, c1, c2):
    table.cell(row, c1).merge(table.cell(row, c2))

def build_table(doc, meta, steps, conclusion):
    rows = len(meta) + 2 + len(steps) + len(conclusion)
    t = doc.add_table(rows=rows, cols=6)
    r = 0
    for k, v in meta:
        merge(t, r, 0, 1); merge(t, r, 2, 5)
        set_cell(t.cell(r, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 2), v)
        r += 1
    merge(t, r, 0, 5)
    set_cell(t.cell(r, 0), '测试过程', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r += 1
    merge(t, r, 1, 2)
    for ci, h in enumerate(['序号', '输入及操作步骤', '', '期望测试结果', '评估准则', '符合性判定\n（√或×或测试结果）']):
        if ci == 2: continue
        set_cell(t.cell(r, ci), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    r += 1
    for num, step_text, expect, criteria, check in steps:
        merge(t, r, 1, 2)
        set_cell(t.cell(r, 0), str(num), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 1), step_text)
        set_cell(t.cell(r, 3), expect)
        set_cell(t.cell(r, 4), criteria, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 5), check, align=WD_ALIGN_PARAGRAPH.CENTER)
        r += 1
    for k, v in conclusion:
        merge(t, r, 0, 1); merge(t, r, 2, 5)
        set_cell(t.cell(r, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(t.cell(r, 2), v)
        r += 1
    doc.add_paragraph()

# Title
for text, align, sz, b in [
    ('5.1.5 功能测试', WD_ALIGN_PARAGRAPH.JUSTIFY, 14, True),
    ('事件目标属性研判功能测试（GN-005）', WD_ALIGN_PARAGRAPH.JUSTIFY, 14, True),
    ('表5-5 GN-005测试用例', WD_ALIGN_PARAGRAPH.CENTER, 12, False),
]:
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.name = 'Times New Roman'
    if b: r.font.bold = True
doc.add_paragraph(); doc.add_paragraph()

build_table(doc, [
    ('测试用例名称', '事件目标属性研判功能测试'),
    ('测试用例标识', 'GN-005'),
    ('测试用例描述', '本测试用例验证系统对事件目标的属性研判能力，覆盖事件的时间属性（日期）、空间属性（地理位置名称及经纬度坐标）、分类属性（事件类型：抗议/冲突/军事活动/其他）、等级属性（特别重大/重大/较大/一般）、状态属性（突发/非突发标记）以及描述属性（事件摘要和关键词分布）。验收人通过点击地图上的事件点和深度分析面板，逐项确认各属性是否完整展示'),
    ('涉及的指标', '指标4研判任务：事件目标属性研判——展示事件的时间、空间、类型、等级、突发状态、摘要等属性'),
    ('前提和约束', '船舶航迹可视化与事件分析系统正常运行，事件数据已加载（地图上可见事件聚类圆/事件点）'),
    ('测试终止条件', '所有测试步骤执行完毕而终止，或因软件运行错误终止'),
], [
    (1, '打开系统页面，点击地图上一个事件聚类圆，在弹出的列表中观察每条事件行', '每行显示：区域标识、事件类型标签（彩色）、日期、突发标记（⚡突发/非突发）、等级标签', '6种属性均可见', ''),
    (2, '点击列表中第一条事件，打开深度分析面板，观察面板顶部类型卡片', '显示事件类型图标+名称（如💥冲突）、等级标签（颜色区分等级）、突发标记、位置名称', '类型、等级、突发、位置4属性醒目展示', ''),
    (3, '在分析面板中查看"事件摘要"和"关键词权重"区域', '事件摘要完整展示描述文本（时间属性隐含在上下文中），关键词条形图展示各关键词权重', '摘要+关键词属性完整', ''),
    (4, '观察地图上该事件点的空间位置，确认与面板中"位置"字段一致', '地图上事件点经纬度与面板显示的位置名称对应（如地图标记在kyiv，面板显示📍 kyiv）', '空间属性一致性', ''),
    (5, '按ESC关闭面板，点击不同区域的事件聚类圆（香港→伊朗→乌克兰），各选一条事件查看属性', '不同区域事件的时间分布、类型倾向、关键词分布均有差异，属性面板正确刷新', '跨区域属性可比对', ''),
    (6, '连续查看3条事件，记录各自的事件类型、等级、突发状态', '3条事件中至少2条的类型或等级不同，体现出系统对事件属性的差异化研判', '属性具有区分度', ''),
], [
    ('测试结论', '  □ 通过      □ 未通过     □ 未测试'),
    ('操作人员', ''),
    ('测试时间', ''),
])

doc.save(r'd:\华东师大计科\6月事件文件\新系统\功能测试表例子_事件研判_v2.docx')
print('Saved GN-005 to v2.docx')
