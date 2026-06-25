# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

TEMPLATE_PATH = r'd:\华东师大计科\6月事件文件\新系统\file_source\功能测试表例子.docx'
OUT_PATH = r'd:\华东师大计科\6月事件文件\新系统\file_source\功能测试表_事件分析KG.docx'

doc = Document(TEMPLATE_PATH)
template_tbl_xml = doc.tables[0]._tbl


def clear_cell(cell):
    for p in cell.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            p.remove(r)
        for fld in p.findall(qn('w:fldSimple')):
            p.remove(fld)
        for hyperlink in p.findall(qn('w:hyperlink')):
            p.remove(hyperlink)


def add_text_to_cell(cell, text):
    first_p = cell.find(qn('w:p'))
    if first_p is None:
        first_p = OxmlElement('w:p')
        cell.insert(0, first_p)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    first_p.append(run)


def set_cell(cell, text):
    clear_cell(cell)
    add_text_to_cell(cell, text)


def unmerge_vertical(tbl_xml):
    for tc in tbl_xml.iter(qn('w:tc')):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            vm = tcPr.find(qn('w:vMerge'))
            if vm is not None:
                tcPr.remove(vm)


case = {
    "gn": "GN-006",
    "name": "\u4e8b\u4ef6\u76ee\u6807\u7ec4\u7ec7\u5173\u8054\u5173\u7cfb\u56fe\u751f\u6210\u6548\u7387",
    "indicator": "\u6307\u68075.1.6\uff1a\u4e8b\u4ef6\u76ee\u6807\u7ec4\u7ec7\u5173\u8054\u5173\u7cfb\u56fe\u751f\u6210<1\u79d2\uff08\u2265100\u76ee\u6807\uff09",
    "module": "\u4e8b\u4ef6\u5206\u6790\u6a21\u5757\uff1a\u77e5\u8bc6\u56fe\u8c31 d3-force \u529b\u5bfc\u5411\u56fe\u6e32\u67d3\uff0c\u653f\u6cbb\u4eba\u7269\u94fe 103 \u4e2a\u8282\u70b9",
    "desc": "\u5728\u4e8b\u4ef6\u5206\u6790\u6a21\u5757\u4e2d\uff0c\u9488\u5bf9 EDAS \u4e8b\u4ef6\u6df1\u5ea6\u5206\u6790\u9762\u677f\u7684\u77e5\u8bc6\u56fe\u8c31\u529f\u80fd\uff0c\u6d4b\u8bd5\u5173\u8054\u5173\u7cfb\u56fe\u7684\u751f\u6210\u6548\u7387\u3002\u7cfb\u7edf\u57fa\u4e8e d3-force \u529b\u5bfc\u5411\u56fe\u5f15\u64ce\uff0c\u9700\u5728 1 \u79d2\u5185\u5b8c\u6210\u5305\u542b 103 \u4e2a\u76ee\u6807\u8282\u70b9\u7684\u77e5\u8bc6\u56fe\u8c31\u6e32\u67d3\uff0c\u5e76\u5728\u9875\u9762\u4e0a\u663e\u793a\u6e32\u67d3\u8017\u65f6\u548c\u8282\u70b9\u6570\u91cf\u3002",
    "premise": "\u667a\u80fd\u60c5\u62a5\u5206\u6790\u5e73\u53f0\u6b63\u5e38\u8fd0\u884c\uff0c\u4e8b\u4ef6\u6570\u636e\u5df2\u52a0\u8f7d\uff0c\u77e5\u8bc6\u56fe\u8c31\u94fe\u6761\u6570\u636e\u6587\u4ef6\uff08chain1.json\u3001chain_kg_1.json\u3001chain_kg.json\uff09\u53ef\u7528",
    "steps": [
        ("\u6253\u5f00\u7cfb\u7edf\u9875\u9762\uff0c\u70b9\u51fb\u9876\u90e8\u5bfc\u822a\u680f\u201c\u4e8b\u4ef6\u5206\u6790\u201d\uff0c\u70b9\u51fb\u5730\u56fe\u4e0a\u4e00\u4e2a\u4e8b\u4ef6\u70b9\uff0c\u5728\u5f39\u7a97\u4e2d\u70b9\u51fb\u201c\u6df1\u5ea6\u5206\u6790\u201d\u6309\u94ae",
         "\u5f39\u51fa EDAS \u6df1\u5ea6\u5206\u6790\u9762\u677f\uff0c\u9ed8\u8ba4\u663e\u793a\u201c\u5206\u6790\u201d\u6807\u7b7e\u9875\uff0c\u9876\u90e8\u6709\u201c\u77e5\u8bc6\u56fe\u8c31\u201d\u548c\u201c\u56e0\u679c\u94fe\u201d\u5207\u6362\u6309\u94ae",
         "\u6df1\u5ea6\u5206\u6790\u9762\u677f\u6b63\u5e38\u5f39\u51fa\uff0c\u4e09\u4e2a\u6807\u7b7e\u9875\u5747\u53ef\u89c1"),
        ("\u70b9\u51fb\u201c\u77e5\u8bc6\u56fe\u8c31\u201d\u6807\u7b7e\u9875\u6309\u94ae",
         "\u9762\u677f\u5207\u6362\u5230\u77e5\u8bc6\u56fe\u8c31\u754c\u9762\uff0c\u663e\u793a\u4e09\u4e2a\u94fe\u6761\u6587\u4ef6\u9009\u62e9\u6309\u94ae\uff1a\u653f\u6cbb\u4eba\u7269\u94fe (Trump\u2192Pence)\u3001\u653f\u6cbb\u4eba\u7269\u94fe (Obama\u2192Biden)\u3001\u536b\u661f\u4efb\u52a1\u94fe (NOAA)\u3002\u9ed8\u8ba4\u81ea\u52a8\u52a0\u8f7d\u7b2c\u4e00\u4e2a\u94fe\u6761",
         "\u77e5\u8bc6\u56fe\u8c31\u754c\u9762\u6b63\u5e38\u52a0\u8f7d\uff0c\u4e09\u4e2a\u94fe\u6761\u6309\u94ae\u53ef\u89c1"),
        ("\u7b49\u5f85\u7b2c\u4e00\u4e2a\u94fe\u6761\uff08\u653f\u6cbb\u4eba\u7269\u94fe Trump\u2192Pence\uff09\u52a0\u8f7d\u5b8c\u6210\uff0c\u67e5\u770b\u77e5\u8bc6\u56fe\u8c31\u5bb9\u5668\u4e0a\u65b9\u7684\u7edf\u8ba1\u4fe1\u606f\u680f",
         "\u7edf\u8ba1\u4fe1\u606f\u680f\u663e\u793a\uff1a\u201c\u6e32\u67d3\u8017\u65f6: \u26a1 XXX ms | \U0001F535 103 \u4e2a\u8282\u70b9\u201d\uff0c\u8017\u65f6\u5b57\u4f53\u4e3a\u91d1\u9ec4\u8272\u5927\u53f7\u52a0\u7c97\uff0c\u8282\u70b9\u6570\u5b57\u4f53\u4e3a\u84dd\u8272\u52a0\u7c97",
         "\u7edf\u8ba1\u4fe1\u606f\u5b8c\u6574\u663e\u793a\uff0c\u8282\u70b9\u6570=103\uff0c\u8017\u65f6\u5b57\u4f53\u589e\u5927\u589e\u4eae"),
        ("\u67e5\u770b SVG \u5bb9\u5668\u5185\u7684\u529b\u5bfc\u5411\u56fe\u6e32\u67d3\u6548\u679c",
         "\u77e5\u8bc6\u56fe\u8c31\u4ee5 d3-force \u529b\u5bfc\u5411\u56fe\u5f62\u5f0f\u5c55\u793a\uff0c\u8282\u70b9\u4ee5\u5706\u5708\u8868\u793a\uff0c\u8282\u70b9\u6807\u7b7e\u663e\u793a\u5728\u5706\u5708\u4e0a\u65b9\u3002\u8fde\u7ebf\u8868\u793a\u8282\u70b9\u4e4b\u95f4\u7684\u5173\u8054\u5173\u7cfb\uff0c\u652f\u6301\u62d6\u62fd\u8282\u70b9\u3001\u6eda\u8f6e\u7f29\u653e\u3001\u60ac\u505c\u9ad8\u4eae",
         "\u529b\u5bfc\u5411\u56fe\u6e32\u67d3\u6210\u529f\uff0c\u4ea4\u4e92\u529f\u80fd\u6b63\u5e38"),
        ("\u70b9\u51fb\u201c\u536b\u661f\u4efb\u52a1\u94fe (NOAA)\u201d\u6309\u94ae\u5207\u6362\u94fe\u6761",
         "\u7edf\u8ba1\u4fe1\u606f\u66f4\u65b0\u4e3a\u201c\U0001F535 44 \u4e2a\u8282\u70b9\u201d\uff0c\u56fe\u5f62\u91cd\u65b0\u6e32\u67d3\u4e3a\u536b\u661f\u4efb\u52a1\u94fe\u7684\u8282\u70b9\u548c\u5173\u7cfb\u3002\u518d\u5207\u6362\u5230\u201c\u653f\u6cbb\u4eba\u7269\u94fe (Obama\u2192Biden)\u201d\uff0c\u663e\u793a\u201c\U0001F535 114 \u4e2a\u8282\u70b9\u201d",
         "\u94fe\u6761\u5207\u6362\u6d41\u7545\uff0c\u8282\u70b9\u6570\u5206\u522b\u663e\u793a 103/114/44"),
    ],
    "remark": "\u77e5\u8bc6\u56fe\u8c31\u57fa\u4e8e d3-force \u529b\u5bfc\u5411\u56fe\u5f15\u64ce\u5b9e\u73b0\uff0c\u4ece fetch \u6570\u636e\u5230\u9996\u6b21 tick \u6e32\u67d3\u5b8c\u6210\u5168\u6d41\u7a0b\u5728 100-600ms \u5185\u3002\u7b2c\u4e00\u4e2a\u94fe\u6761\uff08chain1.json\uff09\u5305\u542b 103 \u4e2a\u8282\u70b9\uff0c\u6ee1\u8db3\u201c\u2265100\u76ee\u6807\u201d\u7684\u6307\u6807\u8981\u6c42\uff0c\u4e14\u6e32\u67d3\u8017\u65f6<1\u79d2\u3002\u8282\u70b9\u6570\u91cf\u548c\u6e32\u67d3\u8017\u65f6\u5728\u77e5\u8bc6\u56fe\u8c31\u5bb9\u5668\u4e0a\u65b9\u4ee5\u5927\u53f7\u4eae\u8272\u5b57\u4f53\u663e\u793a\uff0c\u4fbf\u4e8e\u76f4\u89c2\u9a8c\u8bc1\u3002",
}


def build_table(tbl_xml, case):
    new_tbl = deepcopy(tbl_xml)
    unmerge_vertical(new_tbl)
    rows = new_tbl.findall(qn('w:tr'))

    def rc(row_idx):
        return rows[row_idx].findall(qn('w:tc'))

    set_cell(rc(0)[1], '\u4e8b\u4ef6\u5206\u6790\u2014' + case['name'] + ' \u529f\u80fd\u6d4b\u8bd5')
    set_cell(rc(1)[1], case['gn'])
    set_cell(rc(2)[1], case['desc'])
    set_cell(rc(3)[1], case['indicator'] + '\u2014\u2014' + case['module'])
    set_cell(rc(4)[1], case['premise'])
    set_cell(rc(5)[1], '\u6240\u6709\u6d4b\u8bd5\u6b65\u9aa4\u6267\u884c\u5b8c\u6bd5\u800c\u7ec8\u6b62\uff0c\u6216\u56e0\u8f6f\u4ef6\u8fd0\u884c\u9519\u8bef\u7ec8\u6b62')
    set_cell(rc(6)[0], '\u6d4b\u8bd5\u8fc7\u7a0b')

    for si in range(8):
        ri = 8 + si
        c = rc(ri)
        if si < len(case['steps']):
            inp, exp, crit = case['steps'][si]
            set_cell(c[0], str(si + 1))
            set_cell(c[1], inp)
            set_cell(c[2], exp)
            set_cell(c[3], crit)
            set_cell(c[4], '')
        else:
            for ci in range(len(c)):
                set_cell(c[ci], '')

    set_cell(rc(16)[1], '  \u25a1 \u901a\u8fc7      \u25a1 \u672a\u901a\u8fc7     \u25a1 \u672a\u6d4b\u8bd5')
    return new_tbl


new_doc = Document()
style = new_doc.styles['Normal']
font = style.font
font.name = '\u5b8b\u4f53'
font.size = Pt(10.5)

title_para = new_doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('\u529f\u80fd\u6d4b\u8bd5\u8868\u2014\u2014\u4e8b\u4ef6\u5206\u6790\u77e5\u8bc6\u56fe\u8c31\u90e8\u5206')
title_run.bold = True
title_run.font.size = Pt(16)
title_run.font.name = '\u5b8b\u4f53'

sub_para = new_doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run('\u6307\u68075.1.6 \u4e8b\u4ef6\u76ee\u6807\u7ec4\u7ec7\u5173\u8054\u5173\u7cfb\u56fe\u751f\u6210\u6548\u7387\u529f\u80fd\u6d4b\u8bd5\uff08GN-006\uff09')
sub_run.font.size = Pt(10)
sub_run.font.name = '\u5b8b\u4f53'
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

new_doc.add_paragraph()

tbl_title = new_doc.add_paragraph()
tbl_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tbl_run = tbl_title.add_run(case['gn'] + '  ' + case['name'])
tbl_run.bold = True
tbl_run.font.size = Pt(11)
tbl_run.font.name = '\u5b8b\u4f53'

new_tbl = build_table(template_tbl_xml, case)
new_doc.element.body.append(new_tbl)

new_doc.save(OUT_PATH)
print('Done! Output: ' + OUT_PATH)
